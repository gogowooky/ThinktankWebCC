# -*- coding: utf-8 -*-
"""
walkthrough.mdをWordドキュメントとして出力するスクリプト
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re


def set_cell_shading(cell, color):
    """セルの背景色を設定"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def parse_markdown_table(lines):
    """Markdownテーブルをパース"""
    headers = []
    rows = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith('|'):
            continue
        
        cells = [c.strip() for c in line.split('|')[1:-1]]
        
        if i == 0:
            headers = cells
        elif not all(c.replace('-', '').replace(':', '') == '' for c in cells):
            rows.append(cells)
    
    return headers, rows


def add_table(doc, headers, rows):
    """テーブルを追加"""
    if not headers or not rows:
        return
    
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # ヘッダー行
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        # バッククォートを除去
        cell.text = header.replace('`', '')
        set_cell_shading(cell, '2E74B5')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)
    
    # データ行
    for row_idx, row_data in enumerate(rows, 1):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                # バッククォートを除去
                cell.text = cell_text.replace('`', '')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)


def add_code_block(doc, code_lines, language=''):
    """コードブロックを追加"""
    # コードブロック用の段落
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # 左インデント
        p.paragraph_format.left_indent = Inches(0.3)


def add_mermaid_as_text(doc, mermaid_lines):
    """Mermaid図をテキスト形式で追加"""
    # Mermaid図の説明テキストとして追加
    p = doc.add_paragraph()
    run = p.add_run('[図: Mermaid図 - 元のMarkdownファイルを参照してください]')
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # Mermaidのテキスト内容も追加（参考用）
    for line in mermaid_lines[:15]:  # 最初の15行程度
        p = doc.add_paragraph()
        run = p.add_run('  ' + line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.paragraph_format.space_after = Pt(0)


def convert_markdown_to_docx(md_path, output_path):
    """MarkdownファイルをWordに変換"""
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    code_block_lines = []
    code_language = ''
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # コードブロックの開始/終了
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line.strip()[3:]
                code_block_lines = []
            else:
                in_code_block = False
                if code_language == 'mermaid':
                    add_mermaid_as_text(doc, code_block_lines)
                elif code_language == 'typescript':
                    add_code_block(doc, code_block_lines, code_language)
                else:
                    add_code_block(doc, code_block_lines, code_language)
                code_block_lines = []
                code_language = ''
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # テーブル行の検出
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # テーブル終了
            in_table = False
            headers, rows = parse_markdown_table(table_lines)
            add_table(doc, headers, rows)
            table_lines = []
        
        # 見出し
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
        # 水平線
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        # 空行
        elif line.strip() == '':
            pass  # 空行はスキップ
        # 番号付きリスト
        elif re.match(r'^\d+\.', line.strip()):
            p = doc.add_paragraph()
            # 番号部分と内容を分離
            match = re.match(r'^(\d+\.)\s*\*\*(.+?)\*\*(.*)$', line.strip())
            if match:
                num = match.group(1)
                bold_text = match.group(2)
                rest = match.group(3)
                
                run_num = p.add_run(num + ' ')
                run_num.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                
                run_bold = p.add_run(bold_text)
                run_bold.bold = True
                
                # バッククォートを処理
                rest = rest.replace('`', '')
                run_rest = p.add_run(rest)
            else:
                p.add_run(line.strip())
        # 通常の段落
        else:
            # バッククォートをコードスタイルに変換
            text = line.strip()
            if text:
                p = doc.add_paragraph()
                # バッククォートで囲まれた部分を特別にフォーマット
                parts = re.split(r'(`[^`]+`)', text)
                for part in parts:
                    if part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                    else:
                        # **太字**を処理
                        bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                        for bp in bold_parts:
                            if bp.startswith('**') and bp.endswith('**'):
                                run = p.add_run(bp[2:-2])
                                run.bold = True
                            else:
                                p.add_run(bp)
        
        i += 1
    
    # 残っているテーブルを処理
    if in_table and table_lines:
        headers, rows = parse_markdown_table(table_lines)
        add_table(doc, headers, rows)
    
    # 保存
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Wordファイルを保存しました: {output_path}")


def main():
    md_path = r"C:\Users\gogow\.gemini\antigravity\brain\d77b6335-58db-4ff2-9b02-4856928ebdb7\walkthrough.md"
    output_path = r"C:\Users\gogow\Documents\ThinktankWeb\docs\EventFlowWalkthrough.docx"
    convert_markdown_to_docx(md_path, output_path)


if __name__ == "__main__":
    main()
